import os
from groq import Groq
import re
import unicodedata
import PyPDF2
from docx import Document
from io import BytesIO
import json
import shutil
import time

# -------------------------
# API Key (only yours)
# -------------------------
API_KEY = "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"

# -------------------------
# Setup
# -------------------------
all_results = []
client = Groq(api_key=API_KEY)

# -------------------------
# Helper functions
# -------------------------
def move_failed_resume(file_path, failed_dir):
    """Move resume file to a 'failed_resumes' directory if parsing fails."""
    os.makedirs(failed_dir, exist_ok=True)
    filename = os.path.basename(file_path)
    dest_path = os.path.join(failed_dir, filename)
    shutil.move(file_path, dest_path)
    print(f"Resume Moved {file_path} to {failed_dir}")

def move_processed_resume(file_path, processed_dir):
    """Move successfully processed resume to auto_parsing folder."""
    os.makedirs(processed_dir, exist_ok=True)
    filename = os.path.basename(file_path)
    dest_path = os.path.join(processed_dir, filename)
    shutil.move(file_path, dest_path)
    print(f"✅ Resume moved {file_path} -> {processed_dir}")

def check_pdf_file(file_path):
    _, ext = os.path.splitext(file_path)
    return ext.lower() in [".pdf", ".docx"]

def extract_text_from_pdf(file_bytes):
    try:
        pdf_stream = BytesIO(file_bytes)
        reader = PyPDF2.PdfReader(pdf_stream)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text.strip()
    except Exception as e:
        return f"Error reading PDF: {e}"

def extract_text_from_docx(file_bytes):
    try:
        doc = Document(BytesIO(file_bytes))
        text = []
        for para in doc.paragraphs:
            if para.text.strip():
                text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text.append(cell.text.strip())
        return "\n".join(text).strip()
    except Exception as e:
        return f"Error reading DOCX: {e}"

def clean_and_normalize_text(text):
    text = unicodedata.normalize("NFKC", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n\s*\n+", "\n\n", text)
    text = "\n".join(line.strip() for line in text.splitlines())
    return text.strip()

# -------------------------
# Call LLM directly
# -------------------------
def call_llm(prompt):
    """Send a single request to Groq with your API key."""
    completion = client.chat.completions.create(
        model="openai/gpt-oss-120b",
        messages=[{"role": "user", "content": prompt}],
        temperature=1,
        max_completion_tokens=8192,
        top_p=1,
        reasoning_effort="medium"
    )
    return completion.choices[0].message.content


# -------------------------
# Process all files in folder
# -------------------------
def run_pipeline(json_file, csv_file, folder_path, failed_dir, processed_dir):
    processed_count = 0
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)

        if os.path.isfile(file_path) and check_pdf_file(file_path):
            print(f"Processing: {file_path}")

            # Read file bytes
            with open(file_path, "rb") as f:
                file_bytes = f.read()

            # Extract text
            if filename.lower().endswith(".pdf"):
                extracted_text = extract_text_from_pdf(file_bytes)
            else:
                extracted_text = extract_text_from_docx(file_bytes)

            # Clean text
            cleaned_text = clean_and_normalize_text(extracted_text)

            if not cleaned_text.strip():  # Empty OCR output
                move_failed_resume(file_path, failed_dir)
                continue

            # Prepare LLM prompt
            prompt = f"""
    You are given a resume. Read it carefully and extract the required details exactly as stated — do NOT invent or infer missing data (except for skill expansion per mapping below).

    ### OUTPUT FORMAT
    Return valid JSON only (no comments, no extra text). Use this schema:

    {{
    "personal_information": {{
        "name": "",
        "email": "",
        "phone": "",
        "address": "",
        "other_details": ""
    }},
    "education": [
        {{
        "degree_and_field": "",   # combine degree + field_of_study here
        "institution": "",
        "start_date": "",
        "end_date": ""
        }}
    ],
    "experience": [
        {{
        "job_title": "",
        "company_name": "",
        "job_location": "",
        "start_date": "",
        "end_date": ""
        }}
    ],
    "total_years_experience": "",   # calculate sum of years from all experience
    "skills": {{
        "programming_languages": [],
        "bi_tools": [],
        "version_control_systems": [],
        "testing_and_automation": [],
        "cloud": [],
        "frameworks_libraries": [],
        "databases": [],
        "secondary_skills": []
    }}
    }}

    Resume:
    \"\"\"{cleaned_text}\"\"\"
    """

            # inside your loop:
            output_str = call_llm(prompt)

            try:
                parsed_json = json.loads(output_str)   # validate JSON

                # --- Validation: Check if essential fields are missing ---
                personal_info = parsed_json.get("personal_information", {})
                required_fields = ["name", "email", "phone"]
                missing_fields = [field for field in required_fields if not personal_info.get(field)]

                if missing_fields:
                    print(f"⚠️ Missing {missing_fields} in {filename}. Moving to failed_dir.")
                    move_failed_resume(file_path, failed_dir)
                    continue  # skip appending to results
                all_results.append(parsed_json)        # add to results list
                move_processed_resume(file_path, processed_dir)

            except json.JSONDecodeError:
                print(f"⚠️ Invalid JSON for {filename}. Moving to failed_dir.")
                move_failed_resume(file_path, failed_dir)
                continue

            processed_count += 1
            if processed_count % 7 == 0:
                print("⏸️ Processed 7 files, waiting 60 seconds before making request...")
                time.sleep(60)

    with open(json_file, "w", encoding="utf-8") as f:
        json.dump(all_results, f, indent=2, ensure_ascii=False)

    print(f"✅ All outputs saved in {json_file}")